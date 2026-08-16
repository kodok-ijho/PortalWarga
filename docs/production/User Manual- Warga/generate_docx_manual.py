import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D1DFD7", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def add_callout(doc, text, title="TIPS PINTAR", bg_hex="F0FDF4", border_hex="10B981", icon="💡"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    
    cell = table.cell(0, 0)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=180)
    
    # Left border only
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/>'
        f'<w:top w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    run_title = p.add_run(f"{icon} {title}: ")
    run_title.bold = True
    run_title.font.name = "Segoe UI"
    run_title.font.size = Pt(10)
    run_title.font.color.rgb = RGBColor(18, 62, 43)
    
    run_text = p.add_run(text)
    run_text.font.name = "Segoe UI"
    run_text.font.size = Pt(9.5)
    run_text.font.color.rgb = RGBColor(40, 60, 50)
    
    # Space after table
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

def add_section_header(doc, number, title, badge_text=""):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(5.2)
    table.columns[1].width = Inches(1.3)
    
    cell_left = table.cell(0, 0)
    cell_right = table.cell(0, 1)
    
    set_cell_background(cell_left, "123E2B")
    set_cell_background(cell_right, "123E2B")
    set_cell_margins(cell_left, top=120, bottom=120, left=160, right=100)
    set_cell_margins(cell_right, top=120, bottom=120, left=50, right=140)
    
    # Gold left border
    tcPr = cell_left._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="36" w:space="0" w:color="F6C042"/>'
        f'<w:top w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p_left = cell_left.paragraphs[0]
    p_left.paragraph_format.space_before = Pt(0)
    p_left.paragraph_format.space_after = Pt(0)
    
    run_num = p_left.add_run(f"BAB {number} • ")
    run_num.font.name = "Segoe UI"
    run_num.font.size = Pt(11)
    run_num.font.bold = True
    run_num.font.color.rgb = RGBColor(246, 192, 66)
    
    run_title = p_left.add_run(title)
    run_title.font.name = "Segoe UI"
    run_title.font.size = Pt(11.5)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(255, 255, 255)
    
    if badge_text:
        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_right.paragraph_format.space_before = Pt(0)
        p_right.paragraph_format.space_after = Pt(0)
        run_badge = p_right.add_run(badge_text)
        run_badge.font.name = "Segoe UI"
        run_badge.font.size = Pt(8.5)
        run_badge.font.bold = True
        run_badge.font.color.rgb = RGBColor(246, 192, 66)
    
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(0)
    p_sp.paragraph_format.space_after = Pt(6)

def add_step(doc, step_label, title, description, img_paths_with_captions, base_dir):
    p_step = doc.add_paragraph()
    p_step.paragraph_format.space_before = Pt(8)
    p_step.paragraph_format.space_after = Pt(2)
    
    run_tag = p_step.add_run(f"[{step_label}] ")
    run_tag.font.name = "Segoe UI"
    run_tag.font.size = Pt(10)
    run_tag.font.bold = True
    run_tag.font.color.rgb = RGBColor(16, 185, 129)
    
    run_t = p_step.add_run(title)
    run_t.font.name = "Segoe UI"
    run_t.font.size = Pt(11)
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(18, 62, 43)
    
    p_desc = doc.add_paragraph(description)
    p_desc.paragraph_format.space_before = Pt(2)
    p_desc.paragraph_format.space_after = Pt(6)
    p_desc.runs[0].font.name = "Segoe UI"
    p_desc.runs[0].font.size = Pt(9.5)
    p_desc.runs[0].font.color.rgb = RGBColor(60, 75, 68)
    
    # Add screenshots in a table row
    if img_paths_with_captions:
        cols_count = len(img_paths_with_captions)
        table = doc.add_table(rows=1, cols=cols_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        col_w = Inches(6.5 / cols_count)
        for i in range(cols_count):
            table.columns[i].width = col_w
            
        for i, (img_rel, caption) in enumerate(img_paths_with_captions):
            cell = table.cell(0, i)
            set_cell_background(cell, "F8FAF8")
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            
            full_img_path = os.path.join(base_dir, img_rel)
            if os.path.exists(full_img_path):
                img_width = Inches(5.8 / cols_count)
                p.add_run().add_picture(full_img_path, width=img_width)
            
            p_cap = cell.add_paragraph(caption)
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(2)
            p_cap.paragraph_format.space_after = Pt(0)
            p_cap.runs[0].font.name = "Segoe UI"
            p_cap.runs[0].font.size = Pt(8)
            p_cap.runs[0].font.italic = True
            p_cap.runs[0].font.color.rgb = RGBColor(90, 115, 102)
        
        p_sp = doc.add_paragraph()
        p_sp.paragraph_format.space_before = Pt(0)
        p_sp.paragraph_format.space_after = Pt(6)

def main():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    base_dir = r"D:\DenmasGanteng\Palm Village\Portal Warga Palm Village\PortalPalmVillage\docs\production\User Manual- Warga"
    
    # ================= COVER SECTION =================
    tbl_cover = doc.add_table(rows=1, cols=1)
    tbl_cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_cover.autofit = False
    tbl_cover.columns[0].width = Inches(6.8)
    
    c_cover = tbl_cover.cell(0, 0)
    set_cell_background(c_cover, "123E2B")
    set_cell_margins(c_cover, top=280, bottom=280, left=320, right=320)
    
    p_c1 = c_cover.paragraphs[0]
    p_c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_badge = p_c1.add_run("🌿 PERUMAHAN PALM VILLAGE • BUKU RESMI WARGA")
    r_badge.font.name = "Segoe UI"
    r_badge.font.size = Pt(9.5)
    r_badge.font.bold = True
    r_badge.font.color.rgb = RGBColor(246, 192, 66)
    
    p_c2 = c_cover.add_paragraph()
    p_c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_c2.paragraph_format.space_before = Pt(16)
    p_c2.paragraph_format.space_after = Pt(6)
    r_icon = p_c2.add_run("🌴✨\n")
    r_icon.font.size = Pt(36)
    
    r_t1 = p_c2.add_run("PANDUAN PINTAR\nPORTAL WARGA PALM VILLAGE")
    r_t1.font.name = "Segoe UI"
    r_t1.font.size = Pt(20)
    r_t1.font.bold = True
    r_t1.font.color.rgb = RGBColor(255, 255, 255)
    
    p_c3 = c_cover.add_paragraph()
    p_c3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_c3.paragraph_format.space_before = Pt(6)
    p_c3.paragraph_format.space_after = Pt(14)
    r_sub = p_c3.add_run("Buku Petunjuk Transaksi IPL (QRIS & Transfer Bank), Transparansi Iuran, dan Layanan Hunian")
    r_sub.font.name = "Segoe UI"
    r_sub.font.size = Pt(10.5)
    r_sub.font.color.rgb = RGBColor(209, 231, 220)
    
    p_c4 = c_cover.add_paragraph()
    p_c4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_c4.paragraph_format.space_before = Pt(10)
    r_feat = p_c4.add_run("⚡ QRIS Otomatis Realtime  •  🔑 Login Akun Google Praktis  •  🚨 SOS Satpam Siaga 24 Jam")
    r_feat.font.name = "Segoe UI"
    r_feat.font.size = Pt(8.5)
    r_feat.font.bold = True
    r_feat.font.color.rgb = RGBColor(246, 192, 66)
    
    doc.add_page_break()
    
    # ================= INFOGRAFIS ALUR KILAT =================
    add_callout(
        doc,
        "1. Masuk pakai Google  ➜  2. Cari baris 'RUMAH SAYA' & centang bulan  ➜  3. Scan QRIS di GoPay / BCA / Shopee / Livin  ➜  4. Langsung Lunas Otomatis!",
        title="🚀 ALUR KILAT BAYAR IPL (HANYA 3 MENIT)",
        bg_hex="EEFBF3",
        border_hex="10B981",
        icon="⚡"
    )
    
    # ================= BAB 1: CARA LOGIN =================
    add_section_header(doc, "1", "Cara Masuk / Login ke Portal Warga", "LOGIN GOOGLE")
    add_step(
        doc,
        "Langkah 1.1 - 1.3",
        "Masuk Menggunakan Akun Google Terdaftar",
        "Buka portal warga di browser HP Anda (Chrome/Safari). Klik tombol 'Sign in with Google', pilih akun email Gmail yang sudah terdaftar di pengurus perumahan, lalu konfirmasi sekali klik tanpa repot membuat kata sandi baru.",
        [
            ("screenshots/01-halaman-login.jpg", "1. Klik 'Sign in with Google'"),
            ("screenshots/02-pilih-akun-google.jpg", "2. Pilih Email Terdaftar"),
            ("screenshots/03-konfirmasi-login.jpg", "3. Konfirmasi Masuk Portal")
        ],
        base_dir
    )
    add_callout(
        doc,
        "Jika setelah login nama atau nomor rumah Anda belum muncul, silakan hubungi pengurus RT/RW untuk mengaitkan email Anda dengan data kepemilikan unit.",
        title="BANTUAN AKTIVASI AKUN",
        bg_hex="EFF6FF",
        border_hex="3B82F6",
        icon="ℹ️"
    )
    
    # ================= BAB 2: BERANDA UTAMA =================
    add_section_header(doc, "2", "Mengenal Beranda Utama & Menu Favorit", "BERANDA WARGA")
    add_step(
        doc,
        "Tampilan Beranda",
        "Dashboard Layanan Lengkap dalam Satu Layar",
        "Setelah login, Anda disambut dengan kartu sapaan bertuliskan nama Anda dan badge peran WARGA. Dari beranda ini, Anda dapat mengakses menu Penghuni, Matriks Bayar IPL, serta Kontak Darurat Pos Satpam.",
        [
            ("screenshots/04-beranda-utama.jpg", "Tampilan Beranda Utama Portal")
        ],
        base_dir
    )
    
    # ================= BAB 3: MATRIKS TAGIHAN & RUMAH SAYA =================
    add_section_header(doc, "3", "Melihat Matriks Tagihan & Menemukan 'RUMAH SAYA'", "TRANSPARANSI")
    add_step(
        doc,
        "Langkah 3.1 & 3.2",
        "Filter Tahun Buku & Sorotan Otomatis Unit Rumah Anda",
        "Buka menu 'Matriks Bayar'. Pada bagian atas Anda dapat memilih Tahun Buku (contoh: 2026/2027). Gulir ke bawah untuk menemukan baris unit rumah Anda yang secara otomatis disorot dengan tanda khusus 'RUMAH SAYA'. Anda hanya dapat memilih tagihan pada baris rumah Anda sendiri.",
        [
            ("screenshots/05-matriks-ipl-legenda.jpg", "Legenda Warna & Filter Tahun"),
            ("screenshots/06-tabel-matriks-rumah-saya.jpg", "Sorotan Baris 'RUMAH SAYA'")
        ],
        base_dir
    )
    
    doc.add_page_break()
    
    # ================= BAB 4: PEMBAYARAN QRIS OTOMATIS =================
    add_section_header(doc, "4", "Cara Bayar IPL via QRIS Otomatis (Sangat Direkomendasikan!)", "INSTAN & OTOMATIS")
    add_step(
        doc,
        "Tahap 4.1 - 4.3",
        "Pilih Bulan Tagihan & Lanjut ke QRIS",
        "1. Ketuk kotak bulan yang belum lunas pada baris 'RUMAH SAYA' (pembayaran berurutan jika ada tunggakan).\n2. Panel bawah akan merangkum total rupiah tagihan Anda. Klik tombol kuning 'Lanjutkan Pembayaran →'.\n3. Pada pop-up, pilih metode 'QRIS', lalu klik 'Lanjut ke QRIS'.",
        [
            ("screenshots/07-pilih-bulan-tagihan.jpg", "1. Centang Bulan Tagihan"),
            ("screenshots/08-total-tagihan-lanjut-bayar.jpg", "2. Klik 'Lanjutkan Pembayaran'"),
            ("screenshots/09-pilih-metode-qris.jpg", "3. Pilih Metode QRIS")
        ],
        base_dir
    )
    
    add_step(
        doc,
        "Tahap 4.4 - 4.6",
        "Scan QR Code di E-Wallet & Status Otomatis Hijau",
        "Kode QRIS dinamis unik akan muncul. Buka aplikasi perbankan/e-wallet Anda (GoPay, OVO, ShopeePay, DANA, BCA Mobile, Livin Mandiri, BRImo, Jago, dll.) untuk memindai dan menyelesaikan pembayaran ke merchant 'Palm Village - Social'. Setelah selesai, klik '✓ Saya Sudah Selesai Membayar' dan status tagihan otomatis menjadi Lunas (Hijau) realtime!",
        [
            ("screenshots/10-qr-code-qris.jpg", "4. QR Code QRIS Muncul"),
            ("screenshots/11-bayar-ewallet-m-banking.jpg", "5. Bayar di Aplikasi Bank/E-Wallet"),
            ("screenshots/12-notifikasi-qris-ditutup.jpg", "6. Notifikasi & Sukses Lunas")
        ],
        base_dir
    )
    add_callout(
        doc,
        "Jika hanya memakai 1 HP: Ambil Screenshot tampilan QRIS -> Buka aplikasi e-Wallet/m-Banking -> Pilih menu Scan QR -> Tekan ikon 'Galeri' -> Pilih gambar screenshot tadi -> Langsung bayar!",
        title="TRIK BAYAR DENGAN 1 HP",
        bg_hex="FEF3C7",
        border_hex="D99A19",
        icon="📱"
    )
    
    # ================= BAB 5: TRANSFER BANK MANUAL =================
    add_section_header(doc, "5", "Cara Bayar via Transfer Bank Manual", "TRANSFER MANUAL")
    add_step(
        doc,
        "Langkah 5.1 - 5.3",
        "Unggah Foto / Screenshot Bukti Transfer",
        "Jika memilih transfer manual ke rekening bank pengurus, silakan transfer terlebih dahulu, lalu pilih tab 'Transfer Bank' pada pop-up. Klik 'Choose File' untuk mengunggah foto bukti transfer dari galeri HP (maks 2MB), tambahkan catatan nama pengirim, lalu klik tombol kuning 'Kirim Bukti Transfer'.",
        [
            ("screenshots/13-metode-transfer-bank-manual.jpg", "1. Pilih Tab Transfer Bank"),
            ("screenshots/14-pilih-foto-bukti-transfer.jpg", "2. Pilih Foto dari Galeri"),
            ("screenshots/15-kirim-bukti-transfer.jpg", "3. Klik 'Kirim Bukti Transfer'")
        ],
        base_dir
    )
    add_callout(
        doc,
        "Setelah bukti transfer dikirim, status kotak akan berwarna Oranye (Menunggu Verifikasi) sementara waktu hingga bendahara memeriksa mutasi rekening dan menyetujui status menjadi Lunas (Hijau).",
        title="PROSES VERIFIKASI MANUAL",
        bg_hex="FFFBEB",
        border_hex="F59E0B",
        icon="⏳"
    )
    
    doc.add_page_break()
    
    # ================= BAB 6: KAMUS WARNA STATUS =================
    add_section_header(doc, "6", "Kamus Warna Status Tagihan Matriks IPL", "PANDUAN WARNA")
    
    p_tbl_intro = doc.add_paragraph("Gunakan tabel referensi berikut untuk memahami arti setiap kotak warna pada matriks pembayaran IPL:")
    p_tbl_intro.runs[0].font.name = "Segoe UI"
    p_tbl_intro.runs[0].font.size = Pt(9.5)
    
    table_status = doc.add_table(rows=7, cols=3)
    table_status.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_status.autofit = False
    table_status.columns[0].width = Inches(1.8)
    table_status.columns[1].width = Inches(1.8)
    table_status.columns[2].width = Inches(2.9)
    
    headers = ["Warna Kotak", "Nama Status", "Keterangan Lengkap untuk Warga"]
    for i, h in enumerate(headers):
        cell = table_status.cell(0, i)
        set_cell_background(cell, "123E2B")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.name = "Segoe UI"
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
    
    rows_data = [
        ("🟢 Hijau Muda", "LUNAS", "Tagihan telah berhasil dibayar & terverifikasi. Tertera nominal dan tanggal bayar.", "D1FAE5"),
        ("🟠 Oranye / Kuning Tua", "MENUNGGU VERIFIKASI", "Bukti transfer manual sudah diunggah, menunggu pengecekan mutasi bendahara.", "FFEDD5"),
        ("🟡 Kuning Pucat", "BELUM BAYAR", "Tagihan bulan berjalan atau baru yang belum dilakukan pembayaran.", "FEF3C7"),
        ("🔴 Merah Muda", "TERLAMBAT / DITOLAK", "Melewati jatuh tempo atau bukti transfer manual ditolak pengurus (salah nominal/buram).", "FEE2E2"),
        ("⚪ Abu-abu", "DIBATALKAN", "Transaksi atau tagihan dibatalkan oleh sistem / pengurus.", "F1F5F9"),
        ("⚫ Hijau Gelap", "SEDANG DIPILIH", "Kotak bulan yang sedang Anda tandai/centang untuk dibayar saat ini.", "E2EBE5")
    ]
    
    for row_idx, (w, s, k, bg) in enumerate(rows_data, start=1):
        for col_idx, text in enumerate([w, s, k]):
            cell = table_status.cell(row_idx, col_idx)
            set_cell_background(cell, "FFFFFF" if row_idx % 2 == 0 else "FBFDFB")
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = "Segoe UI"
            r.font.size = Pt(9)
            if col_idx == 1:
                r.font.bold = True
            r.font.color.rgb = RGBColor(20, 40, 30)
            
    set_table_borders(table_status)
    
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(8)
    p_sp.paragraph_format.space_after = Pt(2)
    
    # ================= BAB 7: LAYANAN DARURAT & FAQ =================
    add_section_header(doc, "7", "Tanya Jawab (FAQ) & Kontak Darurat SOS 24 Jam", "FAQ & BANTUAN")
    
    faqs = [
        ("Bisa bayar 6 bulan atau 1 tahun sekaligus?", "Bisa banget! Anda cukup mengetuk beberapa kotak bulan secara berurutan pada baris rumah Anda. Total akan otomatis dihitung dan bisa dibayar dalam 1 kali scan QRIS."),
        ("Mengapa tidak bisa memilih bulan depan jika bulan ini belum lunas?", "Sistem mewajibkan pelunasan berurutan dari bulan terlama agar tidak ada riwayat tunggakan masa lalu yang terlewat."),
        ("Berapa lama verifikasi pembayaran QRIS?", "Sangat cepat! Pembayaran QRIS terverifikasi secara instan (hitungan detik) realtime melalui webhook sistem perbankan."),
        ("Bagaimana jika salah transfer atau salah upload file bukti?", "Tenang, jangan panik! Segera hubungi bendahara/pengurus RT melalui kontak yang tersedia untuk membantu membatalkan atau menyesuaikan data pembayaran Anda.")
    ]
    
    for q, a in faqs:
        p_faq = doc.add_paragraph()
        p_faq.paragraph_format.space_before = Pt(4)
        p_faq.paragraph_format.space_after = Pt(4)
        
        r_q = p_faq.add_run(f"❓ {q}\n")
        r_q.font.name = "Segoe UI"
        r_q.font.size = Pt(9.5)
        r_q.font.bold = True
        r_q.font.color.rgb = RGBColor(18, 62, 43)
        
        r_a = p_faq.add_run(f"👉 {a}")
        r_a.font.name = "Segoe UI"
        r_a.font.size = Pt(9)
        r_a.font.color.rgb = RGBColor(50, 70, 60)
    
    add_callout(
        doc,
        "Untuk situasi darurat keamanan kompleks, tamu mencurigakan, atau bantuan darurat, gunakan tombol Kontak Darurat di beranda portal:\n• 👮 Satpam Riki: 0895-3231-36366 (WhatsApp & Telepon Langsung)\n• 👮 Satpam Roni: 0881-0103-57049 (WhatsApp & Telepon Langsung)",
        title="POS SATPAM SIAGA 24 JAM",
        bg_hex="FFF1F2",
        border_hex="E11D48",
        icon="🚨"
    )
    
    # Save document
    out_docx = os.path.join(base_dir, "Panduan_Penggunaan_Portal_Warga_Palm_Village.docx")
    try:
        doc.save(out_docx)
        print(f"DOCX successfully generated: {out_docx} (Size: {os.path.getsize(out_docx)} bytes)")
    except PermissionError:
        alt_docx = os.path.join(base_dir, "Panduan_Penggunaan_Portal_Warga_Palm_Village_Updated.docx")
        doc.save(alt_docx)
        print(f"Primary file locked. Saved to updated file: {alt_docx} (Size: {os.path.getsize(alt_docx)} bytes)")

if __name__ == "__main__":
    main()
